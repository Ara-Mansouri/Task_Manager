import json
import os
import shutil  # For file handling
from time import sleep
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document

# File names
TASKS_FILE = "tasks.json"
TASKS_SUMMARY_FILE = "tasks_summary.txt"
TASKS_PDF_FILE = "tasks_summary.pdf"
TASKS_DOCX_FILE = "tasks_summary.docx"
TEMP_DOCX_FILE = "tasks_temp.docx"

class TaskManager:
    def __init__(self):
        """Initialize and load tasks from file."""
        self.tasks = self.load_tasks()
        self.removed_tasks = []  # Track removed tasks

    def load_tasks(self):
        """Load tasks from file if it exists, otherwise return an empty structure."""
        if os.path.exists(TASKS_FILE):
            with open(TASKS_FILE, "r", encoding="utf-8") as file:
                return json.load(file)
        return {"pending": [], "completed": []}

    def save_tasks(self):
        """Save current tasks to the JSON file."""
        with open(TASKS_FILE, "w", encoding="utf-8") as file:
            json.dump(self.tasks, file, indent=4)

    def add_task(self, task):
        """Add a new task to the pending list."""
        self.tasks["pending"].append(task)
        self.save_tasks()
        print(f"✅ Task '{task}' added successfully!")

    def remove_task(self, task):
        """Remove a task from the pending list and add it to removed tasks."""
        if task in self.tasks["pending"]:
            self.tasks["pending"].remove(task)
            self.removed_tasks.append(task)
            self.save_tasks()
            print(f"❌ Task '{task}' removed!")
        else:
            print("⚠ Task not found.")

    def complete_task(self, task):
        """Mark a task as completed."""
        if task in self.tasks["pending"]:
            self.tasks["pending"].remove(task)
            self.tasks["completed"].append(task)
            self.save_tasks()
            print(f"✅ Task '{task}' marked as completed!")
        else:
            print("⚠ Task not found.")
    
    def export_tasks_to_txt(self):
        """Save tasks to a text file using UTF-8 encoding and open it in Notepad."""
        with open(TASKS_SUMMARY_FILE, "w", encoding="utf-8") as file:
            file.write("📌 Pending Tasks:\n")
            file.write("\n".join(self.tasks["pending"]) or "(No pending tasks)")
            file.write("\n\n✅ Completed Tasks:\n")
            file.write("\n".join(self.tasks["completed"]) or "(No completed tasks)")
            file.write("\n\n❌ Removed Tasks:\n")
            file.write("\n".join(self.removed_tasks) or "(No removed tasks)")

        print(f"📄 Task summary saved to {TASKS_SUMMARY_FILE}.")
        os.system(f'notepad {TASKS_SUMMARY_FILE}')        

    def export_tasks_to_pdf(self):
        """Save tasks to a PDF file and open it."""
        c = canvas.Canvas(TASKS_PDF_FILE, pagesize=letter)
        width, height = letter
        y_position = height - 40  # Start writing from the top

        c.setFont("Helvetica-Bold", 14)
        c.drawString(30, y_position, "Task Summary Report")
        c.setFont("Helvetica", 12)

        y_position -= 30  # Move down

        # Writing tasks
        sections = [("📌 Pending Tasks:", self.tasks["pending"]),
                    ("✅ Completed Tasks:", self.tasks["completed"]),
                    ("❌ Removed Tasks:", self.removed_tasks)]

        for title, task_list in sections:
            c.setFont("Helvetica-Bold", 12)
            c.drawString(30, y_position, title)
            y_position -= 20
            c.setFont("Helvetica", 12)

            if task_list:
                for task in task_list:
                    c.drawString(50, y_position, f"- {task}")
                    y_position -= 15
            else:
                c.drawString(50, y_position, "(No tasks)")
                y_position -= 15

            y_position -= 15  # Add space between sections

        c.save()
        print(f"📄 Task summary saved to {TASKS_PDF_FILE}.")
        os.system(f'start {TASKS_PDF_FILE}')  # Opens the PDF file
    def export_tasks_to_docx(self):
        """Save tasks to a Word document, allow editing, and update JSON/PDF after saving."""
        # Save the content to a temp file to avoid locking issues
        doc = Document()
        doc.add_heading('Task Summary Report', level=1)

        sections = [("📌 Pending Tasks", self.tasks["pending"]),
                    ("✅ Completed Tasks", self.tasks["completed"]),
                    ("❌ Removed Tasks", self.removed_tasks)]

        for title, task_list in sections:
            doc.add_heading(title, level=2)
            if task_list:
                for task in task_list:
                    doc.add_paragraph(f"- {task}")
            else:
                doc.add_paragraph("(No tasks)")

        doc.save(TEMP_DOCX_FILE)  # Save to a temporary file
        print(f"📄 Task summary saved to {TEMP_DOCX_FILE}. Please edit and save.")

        # Open the document in Word for editing
        os.system(f'start {TEMP_DOCX_FILE}')

        # Wait for the user to finish editing
        input("📝 Press Enter after closing the document in Word...")

        # Try replacing the original file after Word is closed
        try:
            shutil.move(TEMP_DOCX_FILE, TASKS_DOCX_FILE)  # Replace the old file
            print(f"✅ Updated {TASKS_DOCX_FILE} successfully!")
            self.update_tasks_from_docx()  # Update tasks after editing
        except PermissionError:
            print("❌ Error: Word might still be open. Close the file and try again.")

    def update_tasks_from_docx(self):
        """Read tasks from the edited Word file and update the JSON file correctly."""
        doc = Document(TASKS_DOCX_FILE)
        current_section = None
        updated_tasks = {"pending": [], "completed": []}
        removed_tasks = []

        for para in doc.paragraphs:
            text = para.text.strip()

            # Skip empty lines to avoid misclassification
            if not text:
                continue

            # Detect section headers
            if "Pending Tasks" in text:
                current_section = "pending"
                continue
            elif "Completed Tasks" in text:
                current_section = "completed"
                continue
            elif "Removed Tasks" in text:
                current_section = "removed"
                continue

            # Ensure tasks are correctly assigned to their sections
            if current_section :  
                task_text = text[:].strip()  # Remove "- " prefix and clean text

                if current_section == "pending":
                    updated_tasks["pending"].append(task_text)
                elif current_section == "completed":
                    updated_tasks["completed"].append(task_text)
                elif current_section == "removed":
                    removed_tasks.append(task_text)

        # Update tasks and save
        self.tasks = updated_tasks
        self.removed_tasks = removed_tasks
        self.save_tasks()
        self.export_tasks_to_txt()
        self.export_tasks_to_pdf()

        print("✅ Tasks updated successfully from the edited Word document!")
        

def main():
    manager = TaskManager()

    while True:
        print("\n📝 Task Manager")
        print("1. Add Task")
        print("2. Remove Task")
        print("3. Mark Task as Completed")
        print("4. View Tasks")
        print("5. Export Task List to Notepad")  # Export as .txt
        print("6. Export Task List to PDF")  # Export as .pdf
        print("7. Export Task List to Word (DOCX)")  # New option
        print("8. Exit")

        choice = input("Enter your choice (1-8): ").strip()

        if choice == "1":
            task = input("Enter task: ").strip()
            manager.add_task(task)
        elif choice == "2":
            task = input("Enter task to remove: ").strip()
            manager.remove_task(task)
        elif choice == "3":
            task = input("Enter task to mark as completed: ").strip()
            manager.complete_task(task)
        elif choice == "4":
            manager.list_tasks()
        elif choice == "5":
            manager.export_tasks_to_txt()  # Export to Notepad
        elif choice == "6":
            manager.export_tasks_to_pdf()  # Export to PDF
        elif choice == "7":
            manager.export_tasks_to_docx()  # Export to Word
        elif choice == "8":
            print("👋 Exiting Task Manager. See you later!")
            break
        else:
            print("⚠ Invalid choice. Please enter a number between 1-8.")

if __name__ == "__main__":
    main()
